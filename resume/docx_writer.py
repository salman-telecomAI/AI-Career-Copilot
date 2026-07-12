from docx import Document

def save_resume(resume_text: str, output_path: str):
    doc = Document()

    for line in resume_text.split("\n"):
        line = line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        if line.isupper():
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)

    return output_path